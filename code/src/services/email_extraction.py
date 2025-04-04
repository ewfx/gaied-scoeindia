import os
import pytesseract
import io
import fitz  # PyMuPDF for PDF processing
from PIL import Image
from docx import Document
from unstructured.partition.email import partition_email
from bs4 import BeautifulSoup  # Import BeautifulSoup for HTML parsing
import spacy
import email
from email import policy
from email.parser import BytesParser

# Load NLP model
nlp = spacy.load("en_core_web_sm")


def extract_text_from_pdf(pdf_path):
    doc = fitz.open(pdf_path)
    text = ""
    for page in doc:
        text += page.get_text("text") + "\n"
    return text.strip()


def extract_text_from_docx(docx_path):
    doc = Document(docx_path)
    return "\n".join([para.text for para in doc.paragraphs])


def extract_text_from_image(image_bytes):
    try:
        img = Image.open(io.BytesIO(image_bytes))
        text = pytesseract.image_to_string(img)
        return text.strip()
    except Exception as e:
        return f"Error processing image: {str(e)}"


def extract_text_from_eml(eml_path):
    elements = partition_email(filename=eml_path)
    return "\n".join(str(element) for element in elements)


def extract_text_with_ocr(bytes, type):
    text = ""
    doc = fitz.open(stream=bytes, filetype=type)

    for page_num in range(len(doc)):
        for img_index, img in enumerate(doc[page_num].get_images(full=True)):
            xref = img[0]
            base_image = doc.extract_image(xref)
            image_bytes = base_image["image"]
            image = Image.open(io.BytesIO(image_bytes))

            # Extract text from image using Tesseract OCR
            text += pytesseract.image_to_string(image) + "\n"

    return text.strip()


def extract_email_content_with_nlp(file_path):
    ext = file_path.split(".")[-1].lower()
    email_content, attachment_content = "", ""

    if ext == "pdf":
        email_content = extract_text_from_pdf(file_path)
        # attachment_content = extract_text_with_ocr(file_path)
    elif ext == "docx":
        email_content = extract_text_from_docx(file_path)
    elif ext == "eml":
        email_content = extract_text_from_eml(file_path)

    return email_content, attachment_content


def extract_contextual_data(email_content, attachment_content):
    combined_text = email_content + "\n" + attachment_content
    doc = nlp(combined_text)

    extracted_data = {
        "deal_name": "Not Found",
        "amount": "Not Found",
        "expiration_date": "Not Found",
        "subject": "Unknown Subject",
    }

    for ent in doc.ents:
        if ent.label_ == "MONEY":
            extracted_data["amount"] = ent.text
        elif ent.label_ == "DATE":
            extracted_data["expiration_date"] = ent.text
        elif "deal" in ent.text.lower():
            extracted_data["deal_name"] = ent.text

    return extracted_data


def extract_attachments_from_eml(eml_content):
    """
    Extract attachments from an .eml file.

    :param eml_content: The raw content of the .eml file as bytes
    :return: A list of attachments with their filenames and content
    """
    attachments = []

    # Parse the .eml content
    msg = BytesParser(policy=policy.default).parsebytes(eml_content)

    # Iterate through the email parts to find attachments
    for part in msg.iter_attachments():
        filename = part.get_filename()
        if filename:  # If the part has a filename, it's an attachment
            attachments.append(
                {
                    "filename": filename,
                    "content": part.get_payload(
                        decode=True
                    ),  # Decode the attachment content
                }
            )

    return attachments


def extract_pdf_content_with_fitz(pdf_bytes):
    """
    Extract text content from a PDF file using fitz (PyMuPDF).

    :param pdf_bytes: The raw bytes of the PDF file
    :return: Extracted text content from the PDF
    """
    pdf_text = ""
    try:
        # Open the PDF from bytes
        pdf_document = fitz.open(stream=pdf_bytes, filetype="pdf")
        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]
            pdf_text += page.get_text()  # Extract text from the page
        pdf_document.close()
    except Exception as e:
        print(f"Error extracting PDF content with fitz: {e}.\n Trying OCR...")
        try:
            pdf_text += extract_text_with_ocr(pdf_bytes, "pdf")
        except Exception as e2:
            print(f"Error extracting PDF content with OCR: {e2}")

    if pdf_text == "":
        print("Could not extract PDF content.\n Trying OCR...")
        try:
            pdf_text += extract_text_with_ocr(pdf_bytes, "pdf")
        except Exception as e2:
            print(f"Error extracting PDF content with OCR: {e2}")

    return pdf_text


def extract_docx_content(docx_bytes):
    """
    Extract text content from a .docx file using python-docx.

    :param docx_bytes: The raw bytes of the .docx file
    :return: Extracted text content from the .docx file
    """
    docx_text = ""
    try:
        # Open the .docx file from bytes
        docx_file = io.BytesIO(docx_bytes)
        document = Document(docx_file)
        for paragraph in document.paragraphs:
            docx_text += paragraph.text + "\n"
    except Exception as e:
        print(f"Error extracting .docx content: {e}")
    return docx_text


def extract_email_body_and_subject(eml_content):
    """
    Extract the subject and body of an email from an .eml file.

    :param eml_content: The raw content of the .eml file as bytes
    :return: A dictionary containing the subject and body of the email
    """
    # Parse the .eml content
    msg = BytesParser(policy=policy.default).parsebytes(eml_content)

    # Extract the subject
    subject = msg["subject"]

    # Extract the body (plain text or HTML)
    body = ""
    if msg.is_multipart():
        # Iterate through the parts to find the plain text or HTML body
        for part in msg.walk():
            content_type = part.get_content_type()
            if content_type == "text/plain":
                body = part.get_payload(decode=True).decode(part.get_content_charset())
                break
            elif content_type == "text/html":
                html_content = part.get_payload(decode=True).decode(
                    part.get_content_charset()
                )
                # Convert HTML to plain text using BeautifulSoup
                soup = BeautifulSoup(html_content, "html.parser")
                body = soup.get_text()
                break
    else:
        # If the email is not multipart, get the payload directly
        content_type = msg.get_content_type()
        if content_type == "text/plain":
            body = msg.get_payload(decode=True).decode(msg.get_content_charset())
        elif content_type == "text/html":
            html_content = msg.get_payload(decode=True).decode(
                msg.get_content_charset()
            )
            # Convert HTML to plain text using BeautifulSoup
            soup = BeautifulSoup(html_content, "html.parser")
            body = soup.get_text()

    return {"subject": subject, "body": body}
